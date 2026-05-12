import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)


def add_title(document: Document, draft: dict) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(draft.get("title") or "未命名图文")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


def add_paragraph_block(document: Document, block: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.28)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run(block.get("text") or "")


def normalize_image_for_docx(image_path: Path, temp_dir: Path) -> Path:
    normalized_path = temp_dir / f"{image_path.stem}.png"
    with Image.open(image_path) as image:
        rgb_image = image.convert("RGB")
        rgb_image.save(normalized_path, format="PNG")
    return normalized_path


def add_image_block(document: Document, block: dict, temp_dir: Path) -> None:
    image_path = Path(block.get("imagePath") or "")
    picture_added = False

    if image_path.exists():
        try:
            normalized_path = normalize_image_for_docx(image_path, temp_dir)
            picture_paragraph = document.add_paragraph()
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_paragraph.paragraph_format.space_after = Pt(4)
            picture_paragraph.add_run().add_picture(str(normalized_path), width=Inches(5.8))
            picture_added = True
        except Exception:
            picture_added = False

    if not picture_added:
        missing = document.add_paragraph()
        missing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        missing.paragraph_format.space_after = Pt(4)
        placeholder_run = missing.add_run(f"[图片缺失] {image_path}")
        placeholder_run.italic = True
        placeholder_run.font.size = Pt(9)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def export_docx(draft_path: Path, output_path: Path) -> None:
    draft = json.loads(draft_path.read_text(encoding="utf-8"))
    document = Document()
    configure_document(document)
    add_title(document, draft)

    with tempfile.TemporaryDirectory() as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        for block in draft.get("contentBlocks", []):
            block_type = block.get("type")
            if block_type == "paragraph":
                add_paragraph_block(document, block)
            elif block_type == "image":
                add_image_block(document, block, temp_dir)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: export_draft_docx.py <draft_json_path> <output_docx_path>", file=sys.stderr)
        return 1

    draft_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    export_docx(draft_path, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
